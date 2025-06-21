import os
import logging
from typing import Optional
from pkg.projectvar import Projectvar
from pkg.plugins.translator.base_translator import TranslationClient,OllamaTranslator,BaseTranslator,OpenAITranslator
from pkg.server.process.process_translate import set_translate_process_item,update_translate_finalpath,get_glossary_word_list_use
from pkg.database.database import SessionLocal,engine
from pkg.database import models
import time
import openai
import re
from docx import Document
import ollama
from pkg.plugins.translator.utils import language_map
from lxml import etree
# 语言配置
SUPPORTED_LANGUAGES = {
    "English": {
        "code": "en",
        "name": "English",
        "charset": r'[a-zA-Z\s.,!?]',
        "special_chars": r'[!@#$%^&*()_+\-=\[\]{};\'\\:"|,.<>/?]'
    },
    "Chinese": {
        "code": "zh",
        "name": "Simplified Chinese",
        "charset": r'[\u4e00-\u9fff]',
        "special_chars": r'[！@#￥%……&*（）——+【】{}；\'\\："|，。、《》？]'
    },
    "Japanese": {
        "code": "ja",
        "name": "Japanese",
        "charset": r'[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]',
        "special_chars": r'[！＠＃＄％＾＆＊（）ー＋［］｛｝；\'\\：＂｜、。・「」？]'
    },
    "Korean": {
        "code": "ko",
        "name": "Korean",
        "charset": r'[\uAC00-\uD7AF\u1100-\u11FF\u3130-\u318F]',
        "special_chars": r'[！＠＃＄％＾＆＊（）－＋［］｛｝；\'\\：＂｜，。＜＞？]'
    },
    "French": {
        "code": "fr",
        "name": "French",
        "charset": r'[a-zA-Zàâçéèêëîïôûùüÿñæœ\s.,!?]',
        "special_chars": r'[!@#$%^&*()_+\-=\[\]{};\'\\:"|,.<>/?]'
    }
}
DEFAULT_ROLE_PROMPT = """You are a professional translator, proficient in English, Chinese, Japanese, Korean, and French. 
You have expertise in specialized vocabulary across different fields and understand the cultural nuances of each language.
Your translations are accurate, natural, and maintain the original tone and style of the text.

Language-specific guidelines:
- Chinese: Use Simplified Chinese characters and modern standard Mandarin
- Japanese: Use appropriate keigo (honorific language) when context requires
- Korean: Use appropriate honorific forms based on context
- French: Maintain proper gender agreement and formal/informal distinctions
- English: Use appropriate register and maintain natural flow"""
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

gvar = Projectvar()
models.Base.metadata.create_all(bind=engine)
db = SessionLocal()
def clean_translation_result(text):
    """清理翻译结果，统一去除所有模型可能输出的无关提示、解释、说明等，同时保留换行符"""
    # 强力清理大模型常见解释性输出（多行、句中、跨行）
    strong_patterns = [
        r'The translation of [^\n\r""]+?[""]?.+?[""]? into English.*?(is|would be|:).*?(depending on the context.*?used\.|If it refers to.*?would be appropriate\.|would be the better choice\.|\n|$)',
        r'The translation of .+? would be.*?(\n|$)',
        r'If it refers to .*? the translation would be:.*?(\n|$)',
        r'[""]([^""]+)[""] would be appropriate\.',
        r'[""]([^""]+)[""] would be the better choice\.',
        r'"[^"]*" would be appropriate\.',
        r'"[^"]*" would be the better choice\.',
        r'[A-Za-z\s,]*would be appropriate\.',
        r'[A-Za-z\s,]*would be the better choice\.',
        r'The translation of [^\n\r""]+?[""]?.+?[""]? into English would.*?be:.*?("[^"]*"|\'[^"]*\'|\n|$)',
        r'If it refers to .*? the translation would be:.*?("[^"]*"|\'[^"]*\'|\n|$)',
        r'Without additional context.*?not possible to provide.*?translation.*?(\n|$)',
        r'In English, .*? means .*?(\n|$)',
        r'Depending on the context,.*?(\n|$)',
        r'If you provide more context,.*?(\n|$)',
        r'根据上下文.*?翻译.*?可以是.*?(\n|$)',
        r'如果指的是.*?翻译可以是.*?(\n|$)',
        r'没有更多上下文.*?无法提供更具体的翻译.*?(\n|$)',
        r'如果您能提供更多上下文.*?(\n|$)',
        r'很抱歉[，,。\s]*您似乎没有提供要翻译的文本[。.!\s]*请提供原文，我才能进行翻译[。.!\s]*',
        r'很抱歉.*?无法.*?翻译[。.!\s]*',
        r'Sorry[,.!\s]*you did not provide.*?text to translate[。.!\s]*',
        r'Please provide the original text.*?I can translate[。.!\s]*'
    ]
    for pattern in strong_patterns:
        text_before = text
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.DOTALL)
        if text != text_before:
            print(f"[强力清理] Pattern: {pattern}")
    # 修改后的模式，更严格，确保以 "Translating to" 开头，后跟语言名称，然后是冒号
    main_text_pattern = r'^Translating to [A-Za-z\s]+:\s*.*?(?=\n|$)'
    text = re.sub(main_text_pattern, '', text, flags=re.MULTILINE | re.IGNORECASE | re.DOTALL)
    print(f"清理后文本: {text}")
    # 仅压缩超过两个的连续空行，确保保留原始单个空行和换行符
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()
    return text

def preprocess_text(text):
    """文本预处理，保留换行符"""
    # 保护题型标识
    question_types = [
        r'\(Fill-in-the-blank question\)',
        r'\(Multiple Choice\)',
        r'\(True or False\)',
    ]
    
    # 保护换行符
    text = text.replace('\n', '__NEWLINE__')
    
    # 替换为标记
    for i, pattern in enumerate(question_types):
        text = re.sub(pattern, f'__QTYPE_{i}__', text)
    
    # 保护括号内的空格
    text = re.sub(r'\(\s+\)', '__BLANK__', text)
    
    return text

def postprocess_text(text):
    """文本后处理，恢复换行符"""
    # 还原题型标识
    question_types = {
        '__QTYPE_0__': '(Fill-in-the-blank question)',
        '__QTYPE_1__': '(Multiple Choice)',
        '__QTYPE_2__': '(True or False)',
    }
    
    # 还原标记
    for key, value in question_types.items():
        text = text.replace(key, value)
    
    # 还原括号内的空格
    text = text.replace('__BLANK__', '(  )')
    
    # 还原换行符
    text = text.replace('__NEWLINE__', '\n')
    
    return text

def is_same_language(text, target_language):
    """判断文本是否已经是目标语言，优化对混合语言的识别"""
    # 移除空格和标点符号，只保留实际文本
    cleaned_text = re.sub(r'[\s.,!?;:\'\"()\[\]{}<>@#$%^&*_+\-=|\\`~，。！？；：''""（）【】《》]', '', text)
    if not cleaned_text:
        return False
    
    # 目标是英文的情况
    if target_language == "English":
        # 检查是否包含中文字符
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', cleaned_text)
        
        # 如果有中文字符，需要翻译
        if chinese_chars:
            return False
        
        # 如果是纯英文，是否需要翻译取决于上下文
        # 为了避免漏翻，默认返回False
        return False
    
    # 目标是中文的情况
    if target_language == "Chinese":
        # 检查是否包含非中文字符
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', cleaned_text)
        
        # 如果全是中文，则已经是目标语言
        if len(chinese_chars) == len(cleaned_text):
            return True
        
        # 混合文本，默认需要翻译
        return False
    
    # 其他语言情况，默认不是目标语言
    return False

def should_translate(text, target_language):
    """改进的翻译判断函数，优化对混合文本的处理"""
    # 移除空格和标点符号，只保留实际文本
    cleaned_text = re.sub(r'[\s.,!?;:\'\"()\[\]{}<>@#$%^&*_+\-=|\\`~，。！？；：''""（）【】《》、]', '', text)
    if not cleaned_text:
        return False
    
    # 检查是否为纯数字
    if cleaned_text.isdigit():
        return False
    
    # 检查是否为特殊模式（不需要翻译的内容）
    special_patterns = [
        r'^\d+\.\s*$',              # 题号
        r'^[•\-]\s*$',              # 项目符号
        r'^https?://\S+|www\.\S+$', # URL
        r'^[a-zA-Z]:\\[^\\/:*?"<>|\r\n]*$',  # 文件路径
        r'^__[A-Z_]+_\d+__$',       # 特殊标记
    ]
    
    for pattern in special_patterns:
        if re.match(pattern, text.strip()):
            return False
    
    # 检查是否已经是目标语言
    if is_same_language(text, target_language):
        return False
    
    # 中译英情况：只要包含中文就需要翻译
    if target_language == "English":
        has_chinese = bool(re.search(r'[\u4e00-\u9fff]', text))
        has_english = bool(re.search(r'[a-zA-Z]', text))
        
        # 如果包含中文，需要翻译
        if has_chinese:
            # 检查是否为中英混合，这种情况需要特殊处理
            if has_english:
                # 检查中文内容占比，防止只有少量中文被忽略
                chinese_ratio = len(re.findall(r'[\u4e00-\u9fff]', text)) / len(cleaned_text)
                if chinese_ratio < 0.05:  # 如果中文占比很小，可能是英文中的术语解释
                    print(f"混合文本中文占比小于5%，视为主要是英文内容")
                    return True
            return True
            
        # 如果纯英文，这里默认为False，保持原样
        return False
    
    # 英译中情况：包含英文内容就需要翻译
    if target_language == "Chinese":
        has_english = bool(re.search(r'[a-zA-Z]', text))
        has_chinese = bool(re.search(r'[\u4e00-\u9fff]', text))
        
        # 如果已经全是中文，则不需要翻译
        if has_chinese and not has_english:
            return False
        
        # 如果包含英文，检查是否为不需要翻译的专业术语
        if has_english:
            # 检查是否为全大写缩写词或特殊代码
            english_only_terms = [
                r'^[A-Z]{2,}$',  # 全大写缩写词
                r'^[A-Za-z]+\d+$',  # 产品型号等
            ]
            
            if len(text.strip().split()) <= 2:  # 如果只有少量单词
                for pattern in english_only_terms:
                    if re.match(pattern, text.strip()):
                        return False
            
            # 检查英文占比，确保实质内容被翻译
            english_ratio = len(re.findall(r'[a-zA-Z]', text)) / len(cleaned_text)
            if english_ratio < 0.05 and has_chinese:  # 如果英文占比很小且有中文，可能是中文中的术语
                print(f"混合文本英文占比小于5%，视为主要是中文内容")
                return False
                
            return True
    
    # 其他语言情况，默认需要翻译
    return True

def process_text_before_translation(text, target_language):
    """简化版的文本预处理，不做任何内容保护或占位，直接返回原始文本"""
    # 直接返回原始文本，不做任何替换
    return [(text, True)], {}

def find_safe_split_points(text, protected_marks):
    """找到安全的分割点，避免破坏术语完整性"""
    # 获取所有保护标记的位置
    protected_spans = []
    for mark in protected_marks:
        start = 0
        while True:
            start = text.find(mark, start)
            if start == -1:
                break
            protected_spans.append((start, start + len(mark)))
            start += len(mark)
    
    # 找出可能的分割点（句号、问号等）
    potential_splits = []
    for match in re.finditer(r'([。！？!?；;\n])', text):
        potential_splits.append(match.start())
    
    # 筛选不在保护范围内的安全分割点
    safe_splits = []
    for pos in potential_splits:
        is_safe = True
        for start, end in protected_spans:
            if start <= pos < end:
                is_safe = False
                break
        if is_safe:
            safe_splits.append(pos)
    
    return safe_splits

def smart_split_text(text, protected_marks):
    """智能分割文本，保持术语完整性"""
    safe_points = find_safe_split_points(text, protected_marks)
    if not safe_points:
        return [text]  # 找不到安全分割点，返回整个文本
    
    result = []
    last_pos = 0
    for pos in sorted(safe_points):
        result.append(text[last_pos:pos+1])  # 包含分割符
        last_pos = pos+1
    
    if last_pos < len(text):
        result.append(text[last_pos:])
    
    return result

def validate_translation_integrity(original_text, translated_text):
    """验证翻译完整性，确保所有保护内容都正确保留"""
    # 提取所有标记
    original_marks = re.findall(r'__[A-Z_]+_\d+(?:_\d+)?__', original_text)
    translated_marks = re.findall(r'__[A-Z_]+_\d+(?:_\d+)?__', translated_text)
    
    # 数量检查
    if len(original_marks) != len(translated_marks):
        return False, f"标记数量不匹配: 原文 {len(original_marks)} vs 译文 {len(translated_marks)}"
    
    # 内容检查
    missing_marks = []
    for mark in original_marks:
        if mark not in translated_marks:
            missing_marks.append(mark)
    
    if missing_marks:
        return False, f"缺失标记: {missing_marks}"
    
    return True, "翻译完整性验证通过"

def translate_part(text, target_language, is_mixed_language=False):
    """改进的文本翻译函数，针对多语种混合文本优化"""
    print(f"翻译文本: 长度={len(text)}, 传入的is_mixed_language={is_mixed_language}, 目标语言={target_language}")
    if len(text) > 10:
        print(f"文本前50字符: {text[:50]}...")
    # 检查文本是否有效
    if not text or len(text.strip()) <= 1:
        print("文本为空或太短，跳过翻译")
        return text
    # 检查是否全是保护标记
    if re.match(r'^(__[A-Z_]+_\d+(?:_\d+)?__\s*)+$', text):
        print("文本全是保护标记，跳过翻译")
        return text
    lang_config = SUPPORTED_LANGUAGES.get(target_language)
    if not lang_config:
        print(f"不支持的目标语言: {target_language}")
        return text
    # 保存文本中的换行位置
    has_newlines = '\n' in text
    newline_positions = []
    if has_newlines:
        lines = text.split('\n')
        current_pos = 0
        for line in lines[:-1]:
            current_pos += len(line)
            newline_positions.append(current_pos)
            current_pos += 1  # 加上换行符本身的长度
    # 根据内容类型和翻译方向构建优化的提示词
    if is_mixed_language:
        print(f"处理混合文本（使用DEFAULT_ROLE_PROMPT），目标语言: {lang_config['name']}")
        system_prompt = DEFAULT_ROLE_PROMPT 
    else:
        print(f"处理非混合文本，目标语言: {lang_config['name']}")
        system_prompt = f"你是一位专业翻译，精通多种语言。你的任务是将文本翻译为{lang_config['name']}，同时严格遵循以下规则：\n1. 保留所有专业术语、缩写和特殊格式\n2. 保持数字、符号和特殊字符不变\n3. 不要在输出中包含任何翻译说明、注释或元数据\n4. 不要重复原文或添加额外解释，只输出翻译结果"
        if lang_config['code'] == 'zh':
            system_prompt = f"你是一位专业翻译，精通多种语言。你的任务是将文本翻译为{lang_config['name']}，同时严格遵循以下规则：\n1. 保留所有专业术语、缩写和特殊格式。\n2. 保持数字、符号和特殊字符不变。\n3. 不要在输出中包含任何翻译说明、注释或元数据。\n4. 不要重复原文或添加额外解释，只输出翻译结果。"
    prompt = f"将以下文本翻译为{lang_config['name']}：\n\n{text}"
    if lang_config['code'] == 'zh':
        prompt = f"请将以下文本翻译成{lang_config['name']}：\n\n{text}"
    try:
        modelinfo = gvar.get_model_info()
        if modelinfo.get("api_key") == 'ollama':
            client = ollama.Client(host=modelinfo.get('url'))
        else:
            client = openai.OpenAI(
                api_key = modelinfo.get('api_key'),
                base_url= modelinfo.get('url'),
            )
        print("调用API进行翻译...")
        response = client.chat.completions.create(
            model=modelinfo.get('model_selected'),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=4096,
            top_p=1.0,
            frequency_penalty=0.0,
            presence_penalty=0.0
        )
        translated_text = re.sub(r"^<think>.+?</think>","",response.choices[0].message.content.strip(), count=1, flags=re.DOTALL).strip()
        if len(translated_text) > 10:
            print(f"原始翻译结果前50字符: {translated_text[:50]}...")
        # 清理翻译结果
        result = clean_translation_result(translated_text)
        if len(result) > 10:
            print(f"清理后翻译结果前50字符: {result[:50]}...")
        if not result.strip():
            print("警告: 清理后的翻译结果为空，返回原文")
            return text
        # 如果原文有换行，但翻译结果没有足够的换行，尝试恢复换行位置
        if has_newlines and newline_positions and '\n' not in result:
            print("恢复原文换行位置...")
            result_len = len(result)
            text_len = len(text)
            new_result = ""
            last_pos = 0
            for pos in newline_positions:
                relative_pos = min(result_len - 1, int(pos * result_len / text_len))
                best_pos = relative_pos
                for i in range(relative_pos, min(relative_pos + 10, result_len)):
                    if i >= result_len:
                        break
                    if result[i] in " ,.!?;:，。！？；：":
                        best_pos = i + 1
                        break
                new_result += result[last_pos:best_pos] + "\n"
                last_pos = best_pos
            if last_pos < result_len:
                new_result += result[last_pos:]
            result = new_result
        original_length = len(text.strip())
        translated_length = len(result.strip())
        if original_length > 20 and translated_length < original_length * 0.5:
            print(f"警告: 翻译结果明显短于原文 ({translated_length}/{original_length})，可能丢失了内容")
        # 单轮漏翻补漏机制
        # 检测未翻译片段（以目标语言为准）
        if lang_config['code'] == 'zh':
            untranslated = set(re.findall(r'[a-zA-Z0-9_\-]+', result))
        else:
            untranslated = set(re.findall(r'[\u4e00-\u9fff]+', result))
        filtered_untranslated = set()
        for phrase in untranslated:
            # 跳过术语表标志 __XXX__
            if re.fullmatch(r'__[^_]+_\d+__', phrase):
                continue
            # 跳过特殊占位符 __换行符__ 等
            if re.fullmatch(r'__.*?__', phrase):
                continue
            # 跳过大写/混合英文单词（专有名词/品牌/缩写）
            if re.fullmatch(r'[A-Z][a-zA-Z0-9_\-]*', phrase):
                continue
            # 跳过与目标语言一致的内容
            if lang_config['code'] == 'zh' and re.fullmatch(r'[\u4e00-\u9fff]+', phrase):
                continue
            if lang_config['code'] != 'zh' and re.fullmatch(r'[a-zA-Z]+', phrase):
                continue
            filtered_untranslated.add(phrase)
        for phrase in filtered_untranslated:
            # 使用极简直译提示词
            if lang_config['code'] == 'zh':
                sub_prompt = f"只将下列内容翻译成中文，直接输出翻译结果，不要添加任何说明或解释：\n{phrase}"
            else:
                sub_prompt = f"只将下列内容翻译成{lang_config['name']}，直接输出翻译结果，不要添加任何说明或解释：\n{phrase}"
            sub_response = client.chat.completions.create(
                model=modelinfo.get('model_selected'),
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": sub_prompt}
                ],
                temperature=0.1,
                max_tokens=1024,
                top_p=1.0,
                frequency_penalty=0.0,
                presence_penalty=0.0
            )
            sub_trans = re.sub(r"^<think>.+?</think>","",sub_response.choices[0].message.content.strip(), count=1, flags=re.DOTALL).strip()
            sub_trans = clean_translation_result(sub_trans)
            if sub_trans and sub_trans != phrase:
                print(f"补漏翻译: '{phrase}' -> '{sub_trans}'")
                result = result.replace(phrase, sub_trans)
            else:
                print(f"补漏未变: '{phrase}' 保留原文")
        return result
    except Exception as e:
        print(f"翻译出错 ({target_language}): {str(e)}")
        return text

def optimize_batch_division(parts, batch_size=5):
    """优化批次划分，确保术语不跨批次"""
    if not parts:
        return []
        
    optimized_batches = []
    current_batch = []
    current_batch_text = ""
    
    for part, should_translate in parts:
        # 检查当前部分是否包含保护标记
        has_term = bool(re.search(r'__[A-Z_]+_\d+(?:_\d+)?__', part))
        
        # 如果批次为空，直接添加
        if not current_batch:
            current_batch.append((part, should_translate))
            current_batch_text = part
            continue
        
        # 如果当前部分或已有批次包含术语，检查是否需要开始新批次
        if has_term or bool(re.search(r'__[A-Z_]+_\d+(?:_\d+)?__', current_batch_text)):
            # 检查是否有未闭合的术语标记
            all_marks = re.findall(r'__[A-Z_]+_\d+(?:_\d+)?__', current_batch_text + part)
            mark_counts = {}
            for mark in all_marks:
                mark_counts[mark] = mark_counts.get(mark, 0) + 1
            
            # 如果所有标记都成对出现（可能是误报），或批次已满，开始新批次
            all_paired = all(count % 2 == 0 for count in mark_counts.values())
            
            if (len(current_batch) >= batch_size and all_paired) or len(current_batch) >= batch_size * 2:
                optimized_batches.append(current_batch)
                current_batch = [(part, should_translate)]
                current_batch_text = part
            else:
                current_batch.append((part, should_translate))
                current_batch_text += part
        else:
            # 普通文本，如果批次已满则开始新批次
            if len(current_batch) >= batch_size:
                optimized_batches.append(current_batch)
                current_batch = [(part, should_translate)]
                current_batch_text = part
            else:
                current_batch.append((part, should_translate))
                current_batch_text += part
    
    # 添加最后一个批次
    if current_batch:
        optimized_batches.append(current_batch)
    
    return optimized_batches

def restore_protected_content(text, protected_contents):
    """改进的保护内容还原函数，考虑优先级"""
    result = text
    
    # 首先按键的长度排序，优先处理较长的键（通常是更具体的内容）
    sorted_keys = sorted(protected_contents.keys(), key=len, reverse=True)
    
    # 先处理高优先级术语
    high_priority_keys = [k for k in sorted_keys if any(tag in k for tag in 
                         ['TERM_ML', 'TERM_ML_COMMA', 'TERM_WITH_EXPLANATION', 'ENG_WITH_CN_EXPLANATION', 'NORM_TERM'])]
    for key in high_priority_keys:
        result = result.replace(key, protected_contents[key])
    
    # 然后处理其他保护内容
    normal_keys = [k for k in sorted_keys if k not in high_priority_keys]
    for key in normal_keys:
        result = result.replace(key, protected_contents[key])
    
    return result

def translate_text(text, base_lang, target_language):
    """改进的翻译函数，针对多语种混合文本优化，增加内容保护机制"""
    if not text or len(text.strip()) <= 1:
        return text
    
    # 基本的文本预处理，只保留换行符等基本格式
    text = preprocess_text(text)
    
    # 按段落分割文本
    paragraphs = text.split('\n')
    translated_paragraphs = []
    
    # 术语表预处理（如果有）
    from pkg.plugins.translator.utils import preprocess_terms, postprocess_terms, term_mappings
    glossary = gvar.get_glossary()
    
    for paragraph in paragraphs:
        if not paragraph.strip():
            # 保留空行
            translated_paragraphs.append(paragraph)
            continue
            
        # 术语表处理
        if glossary != -1:
            glossary_word = get_glossary_word_list_use(db, glossary, language_map.get(base_lang), language_map.get(target_language))
            processed_paragraph, _, is_direct_translation, direct_translation_content = preprocess_terms(paragraph, glossary_word)
            
            if is_direct_translation:
                translated_paragraphs.append(direct_translation_content)
                continue
        else:
            processed_paragraph = paragraph
        
        # 检测是否需要翻译
        if should_translate(processed_paragraph, target_language):
            # 内容保护处理 - 新增的步骤
            parts, protected_contents = process_text_before_translation(processed_paragraph, target_language)
            processed_text = parts[0][0] if parts else processed_paragraph
            
            # 检测是否为中英混合文本
            has_chinese = bool(re.search(r'[\u4e00-\u9fff]', processed_text))
            has_english = bool(re.search(r'[a-zA-Z]', processed_text))
            is_mixed = has_chinese and has_english
            
            # 调用translate_part进行翻译
            translated = translate_part(processed_text, target_language, is_mixed_language=is_mixed)
            
            # 恢复被保护的内容 - 新增的步骤
            if protected_contents:
                translated = restore_protected_content(translated, protected_contents)
            
            # 后期处理
            paragraph_result = postprocess_text(translated)
            paragraph_result = clean_translation_result(paragraph_result)
            
            # 术语后处理（如果有）
            if glossary != -1:
                paragraph_result = postprocess_terms(paragraph_result, term_mappings)
                
            translated_paragraphs.append(paragraph_result)
        else:
            # 不需要翻译的段落直接保留
            translated_paragraphs.append(processed_paragraph)
    
    # 使用换行符连接翻译后的段落，保留原始段落结构
    result = '\n'.join(translated_paragraphs)
    return result

def paragraph_contains_image(paragraph):
    """检测段落是否包含图片"""
    try:
        # 使用细化的xpath查询，不使用namespaces参数
        return bool(
            paragraph._p.xpath(".//w:drawing//pic:pic")
        )
    except Exception as e:
        print(f"检测图片时出错：{str(e)}")
        # 尝试更简单的查询
        try:
            return bool(paragraph._p.xpath(".//w:drawing"))
        except Exception as e2:
            print(f"备用检测图片时出错：{str(e2)}")
            return False

def run_contains_image(run):
    """检测run是否包含图片"""
    try:
        # 使用细化的xpath查询，不使用namespaces参数
        return bool(
            run._r.xpath("./w:drawing//pic:pic")
        )
    except Exception as e:
        print(f"检测run图片时出错：{str(e)}")
        # 尝试更简单的查询
        try:
            return bool(run._r.xpath("./w:drawing"))
        except Exception as e2:
            print(f"备用检测run图片时出错：{str(e2)}")
            return False

def run_contains_text(run):
    """检测run是否包含文本"""
    try:
        # 使用标准的xpath查询，不使用namespaces参数
        return bool(
            run._r.xpath("./w:t")
        )
    except Exception as e:
        print(f"检测run文本时出错：{str(e)}")
        # 检查常规文本属性
        try:
            return bool(run.text)
        except Exception as e2:
            print(f"备用检测run文本时出错：{str(e2)}")
            return False

def run_contains_text_and_image(run):
    """检测run是否同时包含文本和图片"""
    return run_contains_text(run) and run_contains_image(run)

def update_run_text_preserve_image(run, new_text):
    """更新run的文本内容，但保留图片元素"""
    try:
        # 找到所有文本元素
        text_elements = run._r.xpath("./w:t")
        
        # 如果有文本元素，更新第一个，删除其余的
        if text_elements:
            text_elements[0].text = new_text
            for elem in text_elements[1:]:
                elem.getparent().remove(elem)
        else:
            # 如果没有文本元素，但需要添加文本，在drawing元素前添加文本元素
            drawing_element = run._r.xpath("./w:drawing")[0]
            
            # 创建新的文本元素
            new_t_element = etree.Element("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
            new_t_element.text = new_text
            
            # 添加到drawing元素前
            run._r.insert(run._r.index(drawing_element), new_t_element)
        
        return True
    except Exception as e:
        print(f"更新run文本保留图片时出错：{str(e)}")
        # 如果上述方法失败，尝试简单地设置run.text
        try:
            run.text = new_text
            return True
        except Exception as e2:
            print(f"备用更新run文本时出错：{str(e2)}")
        return False

def process_run_formatting(run):
    """处理 run 的格式"""
    formatting = {
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'font': run.font.name,
        'size': run.font.size,
        'color': run.font.color.rgb if run.font.color else None,
        'hyperlink': run.hyperlink if hasattr(run, 'hyperlink') else None,
        'bookmark': run.bookmark if hasattr(run, 'bookmark') else None,
        'has_image': run_contains_image(run),
        'has_text': run_contains_text(run),
        'has_text_and_image': run_contains_text_and_image(run)
    }
    return formatting

def apply_formatting(run, formatting):
    """应用格式到 run"""
    run.bold = formatting['bold']
    run.italic = formatting['italic']
    run.underline = formatting['underline']
    if formatting['font']:
        run.font.name = formatting['font']
    if formatting['size']:
        run.font.size = formatting['size']
    if formatting['color']:
        run.font.color.rgb = formatting['color']

def preserve_paragraph_formatting(paragraph):
    """保留段落格式"""
    formatting = {
        'alignment': paragraph.alignment,
        'style': paragraph.style,
        'indentation': {
            'left': paragraph.paragraph_format.left_indent,
            'right': paragraph.paragraph_format.right_indent,
            'first_line': paragraph.paragraph_format.first_line_indent
        },
        'spacing': {
            'before': paragraph.paragraph_format.space_before,
            'after': paragraph.paragraph_format.space_after,
            'line': paragraph.paragraph_format.line_spacing
        }
    }
    return formatting

def apply_paragraph_formatting(paragraph, formatting):
    """应用段落格式"""
    paragraph.alignment = formatting['alignment']
    paragraph.style = formatting['style']
    paragraph.paragraph_format.left_indent = formatting['indentation']['left']
    paragraph.paragraph_format.right_indent = formatting['indentation']['right']
    paragraph.paragraph_format.first_line_indent = formatting['indentation']['first_line']
    paragraph.paragraph_format.space_before = formatting['spacing']['before']
    paragraph.paragraph_format.space_after = formatting['spacing']['after']
    paragraph.paragraph_format.line_spacing = formatting['spacing']['line']

def process_list_formatting(paragraph):
    """处理列表格式"""
    list_info = {
        'is_list': paragraph.style.name.startswith('List') if paragraph.style else False,
        'list_style': paragraph.style.name if paragraph.style else None,
        'list_number': paragraph._p.xpath('.//w:numId')[0].val if paragraph._p.xpath('.//w:numId') else None
    }
    return list_info

def preserve_list_formatting(paragraph):
    """保留列表格式"""
    if paragraph.style and paragraph.style.name.startswith('List'):
        return {
            'style': paragraph.style,
            'number_format': paragraph._p.xpath('.//w:numFmt')[0].val if paragraph._p.xpath('.//w:numFmt') else None,
            'start_at': paragraph._p.xpath('.//w:start')[0].val if paragraph._p.xpath('.//w:start') else None
        }
    return None

def apply_list_formatting(paragraph, list_formatting):
    """应用列表格式"""
    if list_formatting:
        paragraph.style = list_formatting['style']
        # 设置编号格式
        if list_formatting.get('number_format'):
            num_fmt = paragraph._p.xpath('.//w:numFmt')
            if num_fmt:
                num_fmt[0].val = list_formatting['number_format']
        # 设置起始编号
        if list_formatting.get('start_at'):
            start = paragraph._p.xpath('.//w:start')
            if start:
                start[0].val = list_formatting['start_at']

def extract_list_number(text):
    """提取列表编号"""
    # 匹配数字编号（如 1. 2. 等）
    number_match = re.match(r'^(\d+\.\s*)', text)
    if number_match:
        return number_match.group(1)
    
    # 匹配项目符号（如 • - 等）
    bullet_match = re.match(r'^([•\-]\s*)', text)
    if bullet_match:
        return bullet_match.group(1)
    
    return None

def batch_translate_paragraphs(paragraphs, base_lang, target_language, batch_size=10):
    """批量翻译段落"""
    translated_paragraphs = []
    current_batch = []
    batch_indices = []  # 记录每个批次中段落的索引
    
    for i, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text:
            current_batch.append(text)
            batch_indices.append(i)
            
        if len(current_batch) >= batch_size:
            # 翻译当前批次
            translated_texts = translate_batch(current_batch, base_lang, target_language)
            
            # 将翻译结果分配给对应的段落
            for idx, trans_text in zip(batch_indices, translated_texts):
                while len(translated_paragraphs) <= idx:
                    translated_paragraphs.append("")
                translated_paragraphs[idx] = trans_text
                
            current_batch = []
            batch_indices = []
    
    # 处理剩余的文本
    if current_batch:
        translated_texts = translate_batch(current_batch, base_lang, target_language)
        for idx, trans_text in zip(batch_indices, translated_texts):
            while len(translated_paragraphs) <= idx:
                translated_paragraphs.append("")
            translated_paragraphs[idx] = trans_text
    
    return translated_paragraphs

def translate_batch(texts, base_lang, target_language):
    """翻译一批文本"""
    translated_texts = []
    
    # 对每个文本单独翻译，保持段落结构
    for text in texts:
        translated_text = translate_text(text, base_lang, target_language)
        translated_texts.append(translated_text)
        
    return translated_texts

def detect_formatted_runs(paragraph):
    """检测段落中的格式化内容（颜色、高亮等）和特殊符号（数学表达式、下标等）"""
    formatted_runs = []
    
    for i, run in enumerate(paragraph.runs):
        # 检查是否有特殊格式
        has_special_format = False
        format_type = None
        
        # 检查字体颜色
        if hasattr(run.font, 'color') and run.font.color and run.font.color.rgb:
            # 排除黑色和自动颜色
            if run.font.color.rgb != '000000' and run.font.color.rgb:
                has_special_format = True
                format_type = 'COLOR'
        
        # 检查高亮
        if hasattr(run, 'highlight_color') and run.highlight_color:
            has_special_format = True
            format_type = 'HIGHLIGHT'
            
        # 检查加粗
        if run.bold:
            has_special_format = True
            format_type = format_type or 'BOLD'
            
        # 检查斜体
        if run.italic:
            has_special_format = True
            format_type = format_type or 'ITALIC'
            
        # 检查下划线
        if run.underline:
            has_special_format = True
            format_type = format_type or 'UNDERLINE'
        
        # 扩展特殊符号检测范围 - 增加对数学符号、上下标和其他特殊格式的检测
        if run.text:
            # 检查是否包含下标、上标字符
            if re.search(r'[₁₂₃₄₅₆₇₈₉₀ᵀ⁰¹²³⁴⁵⁶⁷⁸⁹]', run.text):
                has_special_format = True
                format_type = 'SPECIAL_SYMBOL'
                
            # 检查是否包含数学符号和公式片段
            if re.search(r'[∑∏∫∂√∞≈≠≤≥±×÷⋅⊆⊇⊂⊃∈∉∪∩]', run.text):
                has_special_format = True
                format_type = 'MATH_SYMBOL'
                
            # 检查对下划线表示法的变量名或数学符号
            if re.search(r'[a-zA-Z]+_[a-zA-Z\d]+', run.text):
                has_special_format = True
                format_type = 'CODE_SYMBOL'
                
            # 检查对英文专业术语的保护（例如"batch size"等）
            if re.search(r'[A-Z][a-z]+ [A-Z][a-z]+|[A-Z][a-z]+ation|[A-Z][a-z]+ment|[A-Z][a-z]+sis', run.text):
                has_special_format = True
                format_type = 'TERM'
                
            # 检查中英混合的专业术语（如"gamma（缩放因子）"）
            if re.search(r'[a-zA-Z]+[\s]*[\(（][^）\)]+[\)）]', run.text):
                has_special_format = True
                format_type = 'MIXED_TERM'
                
            # 检查包含逗号的术语组合（如"Multilayer Perceptron, MLP"）
            if re.search(r'[A-Z][a-z]+([ \-][A-Z][a-z]+)+,\s*[A-Z]{2,}', run.text):
                has_special_format = True
                format_type = 'TERM_WITH_ABBR'
        
        if has_special_format and run.text.strip():
            formatted_runs.append((i, run.text, format_type))
    
    return formatted_runs

def translate_paragraph(paragraph, base_lang, target_language, fileid):
    """改进后的段落翻译函数，优化混合文本处理和文本分配逻辑"""
    print("翻译段落")
    # 保存原始格式
    paragraph_formatting = preserve_paragraph_formatting(paragraph)
    list_formatting = preserve_list_formatting(paragraph)
    
    # 获取段落全文
    paragraph_text = ""
    run_formats = []
    
    # 收集所有run的文本和格式信息
    for i, run in enumerate(paragraph.runs):
        format_info = process_run_formatting(run)
        
        # 如果run只包含图片(没有文本)，跳过文本收集
        if format_info['has_image'] and not format_info['has_text']:
            run_formats.append((len(run.text), format_info))
            continue
        
        # 正常收集文本
        paragraph_text += run.text
        run_formats.append((len(run.text), format_info))
    
    # 如果段落中没有文本内容，只有图片，则直接返回
    if not paragraph_text.strip():
        return
    
    # 提取列表编号
    list_number = extract_list_number(paragraph_text)
    if list_number:
        paragraph_text = paragraph_text[len(list_number):]
    
    # 检查是否需要停止翻译
    isneedstop = gvar.get_needstop()
    if fileid in isneedstop:
        return
    
    # 检测是否为混合文本
    has_chinese = bool(re.search(r'[\u4e00-\u9fff]', paragraph_text))
    has_english = bool(re.search(r'[a-zA-Z]', paragraph_text))
    is_mixed = has_chinese and has_english
    
    print(f"文本内容: {paragraph_text[:50]}... 是否混合文本: {is_mixed}")
    
    # 直接调用translate_text翻译整个段落
    translated_text = translate_text(paragraph_text, base_lang, target_language)
    print(f"翻译结果: {translated_text[:50]}...")
    
    # 检查翻译完整性
    if not check_translation_completeness(paragraph_text, translated_text, target_language):
        print("翻译结果可能不完整，尝试再次翻译...")
        # 确保混合文本标记为True，更明确的指导翻译
        translated_text = translate_part(paragraph_text, target_language, is_mixed_language=True)
    
    # 重新组合列表编号和翻译后的文本
    if list_number:
        translated_text = list_number + translated_text
    
    # 清空原始段落内容并按格式重建，但保留包含图片的run
    for i, run in enumerate(paragraph.runs):
        format_info = run_formats[i][1] if i < len(run_formats) else None
        # 如果run只包含图片(没有文本)，完全跳过
        if format_info and format_info['has_image'] and not format_info['has_text']:
            continue
        # 如果run同时包含文本和图片，使用特殊处理保留图片
        elif format_info and format_info['has_text_and_image']:
            # 暂时清空，后续再设置文本
            continue
        # 如果run只包含文本，清空其内容
        else:
            run.text = ""
    
    # 重新分配文本到runs，尝试保持原始格式，同时处理图片
    remaining_text = translated_text
    
    # 过滤出可以设置文本的runs（包含只有文本的runs和同时包含文本和图片的runs）
    text_runs = [(i, run) for i, run in enumerate(paragraph.runs) 
                if i < len(run_formats) and 
                (run_formats[i][1]['has_text'] and not run_formats[i][1]['has_image'] or 
                 run_formats[i][1]['has_text_and_image'])]
    
    # 如果没有可以设置文本的run，则无需处理
    if not text_runs:
        return
    
    # 如果只有一个可以设置文本的run，直接设置全部文本
    if len(text_runs) == 1:
        i, run = text_runs[0]
        format_info = run_formats[i][1]
        # 如果run同时包含文本和图片，使用特殊处理
        if format_info['has_text_and_image']:
            update_run_text_preserve_image(run, remaining_text)
        else:
            run.text = remaining_text
            apply_formatting(run, format_info)
    else:
        # 多个run的情况，使用改进的文本分配逻辑
        
        # 保存原始分割点信息
        split_points = []
        current_pos = 0
        
        for i, (length, format_data) in enumerate(run_formats):
            if i < len(paragraph.runs) and format_data['has_text']:
                current_pos += length
                if i < len(text_runs) - 1:  # 除了最后一个run
                    # 只记录文本run之间的分割点
                    for j, (run_index, _) in enumerate(text_runs):
                        if run_index == i:
                            split_points.append(current_pos)
                            break
        
        # 计算翻译后文本的分割点
        translated_split_points = []
        if paragraph_text and split_points:  # 确保有原始文本和分割点
            for orig_pos in split_points:
                # 计算相对位置（百分比）
                relative_pos = min(1.0, orig_pos / len(paragraph_text))
                # 在翻译文本中找到相应位置
                trans_pos = int(relative_pos * len(translated_text))
                
                # 找到附近的自然分割点
                best_pos = find_nearest_split_point(translated_text, trans_pos)
                translated_split_points.append(best_pos)
        
        # 如果找不到合适的分割点，使用均匀分配
        if not translated_split_points and len(text_runs) > 1:
            avg_length = len(translated_text) // len(text_runs)
            for i in range(1, len(text_runs)):
                pos = i * avg_length
                translated_split_points.append(pos)
        
        # 根据分割点分配文本
        start_pos = 0
        for j, (i, run) in enumerate(text_runs):
            if j < len(translated_split_points):
                end_pos = translated_split_points[j]
            else:  # 最后一个run获得所有剩余文本
                end_pos = len(translated_text)
            
            # 确保不超出边界
            end_pos = min(end_pos, len(translated_text))
            
            # 获取当前run的文本部分
            new_text = translated_text[start_pos:end_pos]
            format_info = run_formats[i][1]
            
            # 设置文本，区分处理包含图片的run
            if format_info['has_text_and_image']:
                update_run_text_preserve_image(run, new_text)
            else:
                run.text = new_text
                apply_formatting(run, format_info)
            
            start_pos = end_pos
    
    # 应用段落格式
    apply_paragraph_formatting(paragraph, paragraph_formatting)
    if list_formatting:
        apply_list_formatting(paragraph, list_formatting)

def translate_table(table, base_lang,target_language,fileid):
    print("翻译表格")
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                isneedstop = gvar.get_needstop()
                if fileid in isneedstop:
                    break
                translate_paragraph(paragraph, base_lang,target_language,fileid)

def translate_header_footer(header_footer, base_lang,target_language,fileid):
    print("翻译页眉页脚")
    for paragraph in header_footer.paragraphs:
        isneedstop = gvar.get_needstop()
        if fileid in isneedstop:
            break
        translate_paragraph(paragraph, base_lang,target_language,fileid)

def translate_file(input_file: str, output_file: Optional[str] = None, 
                  translator: Optional[BaseTranslator] = None,
                  base_lang:str = 'Chinese',
                  target_language: str = 'English',
                  fileid:str = ""):
    """
    翻译文档的主入口函数，确保保留原文档的段落换行和格式
    
    Args:
        input_file: 输入文件路径
        output_file: 输出文件路径（可选）
        translator: 翻译客户端（可选）
        target_language: 目标语言（默认为英语）
    
    Returns:
        str: 输出文件路径
    """
    if not output_file:
        file_name, file_ext = os.path.splitext(input_file)
        output_file = f"{file_name}_translated{file_ext}"
    try:
        doc = Document(input_file)
        
        # 翻译正文段落
        localtime = time.time()
        update_translate_finalpath(db=db,fileid=fileid,translated_path="",status=0,translated_time="")
        total = len(doc.paragraphs)+len(doc.tables)+len(doc.sections)
        current_num = -1
        for paragraph in doc.paragraphs:
            isneedstop = gvar.get_needstop()
            if fileid in isneedstop:
                break
            current_num += 1
            set_translate_process_item(db=db,fileid=fileid,process=current_num/total)
            translate_paragraph(paragraph, base_lang,target_language,fileid)
        
        # 翻译表格
        for table in doc.tables:
            isneedstop = gvar.get_needstop()
            if fileid in isneedstop:
                break
            current_num += 1
            set_translate_process_item(db=db,fileid=fileid,process=current_num/total)
            translate_table(table, base_lang,target_language,fileid)
        
        # 翻译页眉页脚
        for section in doc.sections:
            isneedstop = gvar.get_needstop()
            if fileid in isneedstop:
                break
            current_num += 1
            set_translate_process_item(db=db,fileid=fileid,process=current_num/total)
            translate_header_footer(section.header, base_lang,target_language,fileid)
            translate_header_footer(section.footer, base_lang,target_language,fileid)
        
        if fileid in isneedstop:
            update_translate_finalpath(db=db,fileid=fileid,translated_path="",status=-1,translated_time='')
            set_translate_process_item(db,fileid,0)
            gvar.delete_needstop(fileid)
            return  
        
        # 保存文档之前确保所有格式正确
        # 设置文档属性，确保保留格式
        set_translate_process_item(db=db,fileid=fileid,process=1)
        update_translate_finalpath(db=db,fileid=fileid,translated_path=output_file,status=1,translated_time=time.time()-localtime)
        
        # 保存文档
        doc.save(output_file)
        return output_file
    except Exception as e:
        print(f"翻译文档时出错: {str(e)}")
        update_translate_finalpath(db=db,fileid=fileid,translated_path="",status=-1,translated_time='')
        set_translate_process_item(db,fileid,0)
        if fileid in gvar.get_needstop():
            gvar.delete_needstop(fileid)
        return None

def find_nearest_split_point(text, position):
    """在文本中找到最接近指定位置的分割点（空格、标点等）"""
    # 搜索的范围（向前和向后多少字符）
    search_range = 10
    
    # 搜索范围的边界
    start = max(0, position - search_range)
    end = min(len(text), position + search_range)
    
    # 如果位置恰好在分割点上，直接返回
    if position < len(text) and text[position] in " ,.!?;:，。！？；：\n":
        return position
    
    # 向后搜索
    for i in range(position, end):
        if i >= len(text):
            break
        if text[i] in " ,.!?;:，。！？；：\n":
            return i
    
    # 向前搜索
    for i in range(position-1, start-1, -1):
        if i < 0:
            break
        if text[i] in " ,.!?;:，。！？；：\n":
            return i + 1  # 返回标点符号后的位置
    
    # 如果没有找到合适的分割点，返回原始位置
    return position

def check_translation_completeness(original, translated, target_language):
    """检查翻译是否完整，没有遗漏重要内容"""
    # 针对中译英
    if target_language == "English":
        # 检查原文中的中文内容是否被保留/翻译
        chinese_chars = re.findall(r'[\u4e00-\u9fff]+', original)
        if chinese_chars:
            # 保守估计：翻译后长度应该不低于原始中文字符数的1.5倍
            min_expected_length = len(''.join(chinese_chars)) * 1.5
            if len(translated) < min_expected_length:
                print(f"警告: 翻译结果可能不完整。原中文字符数: {len(''.join(chinese_chars))}, 翻译结果长度: {len(translated)}")
                return False
                
    # 针对英译中
    elif target_language == "Chinese":
        # 提取原文中的英文单词
        english_words = re.findall(r'[a-zA-Z]+', original)
        if english_words:
            # 保守估计：翻译后长度不应该远低于原始英文单词数
            min_expected_length = len(english_words) * 0.5
            if len(translated) < min_expected_length:
                print(f"警告: 翻译结果可能不完整。原英文单词数: {len(english_words)}, 翻译结果长度: {len(translated)}")
                return False
    
    return True

# # 批量还原 __PRESERVED_x_x__ 占位符的辅助函数
# def restore_preserved(text, preserved_contents):
#     def replacer(match):
#         key = match.group(0)
#         return preserved_contents.get(key, key)
#     return re.sub(r'__PRESERVED_\d+_\d+__', replacer, text)